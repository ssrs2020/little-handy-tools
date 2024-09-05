import random
import string
from docx import Document
# in order to save password file to .docx, u will need to install python-docx package
# can be achieved by using: pip install python-docx

# Function to generate the password
def generate_password(length, include_lowercase, include_uppercase, include_special, include_numbers):
    characters = ""

    if include_lowercase:
        characters += string.ascii_lowercase
    if include_uppercase:
        characters += string.ascii_uppercase
    if include_special:
        characters += string.punctuation
    if include_numbers:
        characters += string.digits

    if not characters:
        raise ValueError("No character sets selected!")

    return ''.join(random.choice(characters) for _ in range(length))


# Function to save the password to a .txt or .docx file
def save_password_to_file(passwords, filename, file_format):
    if file_format == 'txt':
        with open(f"{filename}.txt", "w") as file:
            for account, password in passwords.items():
                file.write(f"Account: {account}, Password: {password}\n")
        print(f"Passwords saved to {filename}.txt")
    elif file_format == 'docx':
        doc = Document()
        doc.add_heading("Generated Passwords", 0)
        for account, password in passwords.items():
            doc.add_paragraph(f"Account: {account}, Password: {password}")
        doc.save(f"{filename}.docx")
        print(f"Passwords saved to {filename}.docx")
    else:
        print("Invalid file format!")


def main():
    passwords = {}  # Dictionary to store account and password pairs

    while True:
        account_name = input("Enter the account name this password is for: ")

        # Get password length
        while True:
            try:
                length = int(input("Enter password length (8-16): "))
                if 8 <= length <= 16:
                    break
                else:
                    print("Please enter a length between 8 and 16.")
            except ValueError:
                print("Invalid input. Please enter a number.")

        # Get character options
        include_lowercase = input("Include lowercase letters? (y/n): ").lower() == 'y'
        include_uppercase = input("Include uppercase letters? (y/n): ").lower() == 'y'
        include_special = input("Include special characters? (y/n): ").lower() == 'y'
        include_numbers = input("Include numbers? (y/n): ").lower() == 'y'

        try:
            password = generate_password(length, include_lowercase, include_uppercase, include_special, include_numbers)
            print(f"Generated Password: {password}")
            passwords[account_name] = password
        except ValueError as e:
            print(e)
            continue

        # Ask if user wants to generate another password
        another = input("Generate another password for this or another account? (y/n): ").lower()
        if another != 'y':
            break

    # Ask if user wants to save the passwords to a file
    save = input("Would you like to save the passwords to a file? (y/n): ").lower()
    if save == 'y':
        filename = input("Enter the filename (without extension): ")
        file_format = input("Enter file format (txt/docx): ").lower()
        save_password_to_file(passwords, filename, file_format)


if __name__ == "__main__":
    main()
